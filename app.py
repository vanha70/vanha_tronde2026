import streamlit as st
from docx import Document
import io, re, random, zipfile, string

# --- GIAO DIỆN THEO MẪU ---
st.set_page_config(page_title="TNMix Pro - GV Nguyễn Văn Hà", layout="centered")
st.markdown("""
    <style>
    [data-testid="stAppViewContainer"] { background: linear-gradient(180deg, #f3605f 0%, #f9a066 100%); }
    .main-container { background-color: white; border-radius: 30px; padding: 30px; margin-top: 10px; box-shadow: 0 10px 30px rgba(0,0,0,0.2); }
    .logo-badge { background: rgba(255,255,255,0.3); padding: 10px 20px; border-radius: 15px; color: white; font-weight: bold; text-align: center; width: fit-content; margin: auto; }
    .teacher-info { text-align: center; color: white; margin-top: 10px; font-size: 1.1em; }
    div.stButton > button:first-child[kind="primary"] { background: linear-gradient(90deg, #f3605f, #f9a066); color: white; border: none; border-radius: 25px; height: 50px; width: 100%; font-weight: bold; font-size: 18px; }
    .upload-area { border: 2px solid #f3605f; border-radius: 20px; padding: 40px; text-align: center; background-color: #fffafb; }
    </style>
    """, unsafe_allow_html=True)

# --- HÀM SAO CHÉP GIỮ NGUYÊN ĐỊNH DẠNG & CÔNG THỨC ---
def copy_para_safe(source_para, target_doc):
    """Sao chép paragraph sang file mới một cách an toàn để tránh lỗi mở file"""
    new_p = target_doc.add_paragraph()
    new_p.paragraph_format.alignment = source_para.alignment
    for run in source_para.runs:
        new_run = new_p.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        # Nếu có công thức hoặc hình ảnh (không phải text), sao chép phần tử XML nội bộ
        if not run.text:
            new_run._r.append(run._r) 
    return new_p

# --- LOGIC NHẬN DIỆN 3 PHẦN ---
def parse_exam_v4(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    parts = {"I": [], "II": [], "III": []}
    current_part = None
    current_q = []

    for para in doc.paragraphs:
        txt = para.text.strip().upper()
        # Nhận diện tiêu đề phần 
        if "PHẦN I" in txt: current_part = "I"; continue
        elif "PHẦN II" in txt: current_part = "II"; continue
        elif "PHẦN III" in txt: current_part = "III"; continue

        if current_part:
            # Nhận diện câu dựa trên nội dung (tự động gom nhóm nếu paragraph không có A. B. C. D.)
            if current_part == "I" and (re.search(r'^[A-D][\.\)]', para.text.strip()) or not para.text.strip()):
                if current_q: current_q.append(para)
            elif current_part == "II" and re.match(r'^[a-d]\)', para.text.strip()):
                if current_q: current_q.append(para)
            else:
                if current_q: parts[current_part].append(current_q)
                current_q = [para]
    
    if current_q: parts[current_part].append(current_q)
    return parts

# --- TẠO ĐỀ VÀ ĐÁP ÁN ---
def generate_exam(parts, code):
    doc = Document()
    doc.add_heading(f"MÃ ĐỀ: {code}", 0)
    ans_key = []

    for p_label, p_key in [("PHẦN I: Trắc nghiệm nhiều lựa chọn", "I"), 
                           ("PHẦN II: Trắc nghiệm đúng sai", "II"), 
                           ("PHẦN III: Trắc nghiệm trả lời ngắn", "III")]:
        if not parts[p_key]: continue
        doc.add_heading(p_label, level=1)
        
        # Trộn câu hỏi
        qs = list(parts[p_key])
        random.shuffle(qs)

        for i, q_paras in enumerate(qs, 1):
            # Paragraph đầu tiên (Thân câu hỏi) [cite: 3, 5, 7]
            p0 = doc.add_paragraph()
            p0.add_run(f"Câu {i}: ").bold = True
            body = re.sub(r'^(Câu|Câu hỏi)\s+\d+[:.]', '', q_paras[0].text, flags=re.I).strip()
            p0.add_run(body)

            # Các paragraph còn lại (Đáp án, hình ảnh) [cite: 4, 6, 8]
            for p in q_paras[1:]:
                # Tự động lấy đáp án đúng nếu có gạch chân
                if p_key == "I":
                    for run in p.runs:
                        if run.underline and re.match(r'^[A-D]', run.text.strip()):
                            ans_key.append(f"C{i}-{run.text.strip()[0]}")
                
                # Sao chép an toàn để giữ công thức/hình ảnh [cite: 76, 107, 113]
                copy_para_safe(p, doc)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf, ans_key

# --- GIAO DIỆN ---
st.markdown('<div class="logo-badge">TNMix</div>', unsafe_allow_html=True)
st.markdown("<h2 style='text-align:center; color:white; margin-bottom:0;'>TNMix Pro - Nguyễn Văn Hà</h2>", unsafe_allow_html=True)
st.markdown(f'<div class="teacher-info">Zalo: 0907781595</div>', unsafe_allow_html=True)

with st.container():
    st.markdown('<div class="main-container">', unsafe_allow_html=True)
    uploaded = st.file_uploader("Chọn file .docx", type=["docx"], label_visibility="collapsed")
    
    if uploaded:
        num = st.number_input("Số mã đề:", 1, 10, 4)
        if st.button("BẮT ĐẦU TRỘN ĐỀ", type="primary"):
            parts = parse_exam_v4(uploaded.read())
            
            if not any(parts.values()):
                st.error("Dữ liệu trống! Hãy kiểm tra từ khóa 'PHẦN I' trong file của thầy.")
            else:
                zip_buf = io.BytesIO()
                with zipfile.ZipFile(zip_buf, "a") as zf:
                    for i in range(num):
                        c = 1201 + i
                        d_buf, k = generate_exam(parts, c)
                        zf.writestr(f"De_{c}.docx", d_buf.getvalue())
                        # Tạo file đáp án riêng cho mỗi mã đề
                        k_txt = f"DAP AN MA DE {c}:\n" + ", ".join(k)
                        zf.writestr(f"DapAn_{c}.txt", k_txt.encode('utf-8'))
                
                st.success("Trộn đề thành công! Công thức và hình ảnh đã được xử lý.")
                st.download_button("📥 TẢI TRỌN BỘ (.ZIP)", zip_buf.getvalue(), "KetQua_TNMix_V4.zip")
    st.markdown('</div>', unsafe_allow_html=True)
